import docx
import threading
from googletrans import Translator
from serbian_text_converter import SerbianTextConverter


class WordTranslationApp:
    def __init__(self, input_path, output_path, target_lang="en"):
        # Initialize paths and target language
        self.input_path = input_path
        self.output_path = output_path
        self.target_language_code = target_lang  # Default to "en" (English)

        # Start the translation process
        self.translate_document()

    def translate_paragraph(self, paragraph, translator, target_lang='en'):
        print(f"Processing paragraph: {paragraph.text}")  # Log the paragraph text being processed
        translated_runs = []
        current_run_text = ''  # This will hold the accumulated text from multiple runs

        # Iterate through all the runs in the paragraph
        for run in paragraph.runs:
            if run.text.strip():  # If there's any text to translate
                print(f"Accumulating text: {run.text}")  # Debugging output to see the accumulated text
                current_run_text += run.text  # Collect the text from the current run
            else:
                # Add the non-translated run (e.g., spaces or line breaks)
                translated_runs.append(run)

        # If any text was accumulated, translate it as a whole
        if current_run_text:
            try:
                # Translate the accumulated text
                translated_text = translator.translate(current_run_text, dest=target_lang)
                print(f"Translated text: {translated_text.text}")

                # Create a new run for the translated text and preserve the original formatting
                translated_run = paragraph.add_run(translated_text.text)

                # Preserve formatting (bold, italic, underline, font size, color, etc.)
                translated_run.bold = run.bold
                translated_run.italic = run.italic
                translated_run.underline = run.underline
                translated_run.font.size = run.font.size
                translated_run.font.color.rgb = run.font.color.rgb
                translated_run.font.name = run.font.name
                translated_run.font.highlight_color = run.font.highlight_color

                # If translating to Serbian Latin, convert Cyrillic to Latin
                if target_lang == "sr_Latn":
                    translated_run.text = SerbianTextConverter.to_latin(translated_run.text)

                # Ensure space is added if necessary after translated text
                if translated_run.text and not translated_run.text.endswith(' '):
                    translated_run.text += ' '

                translated_runs.append(translated_run)
            except Exception as e:
                print(f"Error translating text: {current_run_text}, Error: {e}")

        # Clear the paragraph and rebuild it with translated runs
        paragraph.clear()

        # Re-add the translated runs to the paragraph
        for translated_run in translated_runs:
            paragraph.add_run(translated_run.text).bold = translated_run.bold
            paragraph.runs[-1].italic = translated_run.italic
            paragraph.runs[-1].underline = translated_run.underline
            paragraph.runs[-1].font.size = translated_run.font.size
            paragraph.runs[-1].font.color.rgb = translated_run.font.color.rgb
            paragraph.runs[-1].font.name = translated_run.font.name
            paragraph.runs[-1].font.highlight_color = translated_run.font.highlight_color

        return paragraph

    def process_paragraphs(self, doc, translator, target_lang='en'):
        # Create a list of threads for parallel processing
        threads = []

        # Process all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only process non-empty paragraphs
                thread = threading.Thread(target=self.translate_paragraph, args=(paragraph, translator, target_lang))
                threads.append(thread)
                thread.start()

        # Process tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Translate the text in each cell
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():  # Only process non-empty paragraphs in cells
                            thread = threading.Thread(target=self.translate_paragraph, args=(paragraph, translator, target_lang))
                            threads.append(thread)
                            thread.start()

        # Process headers and footers (if any)
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        thread = threading.Thread(target=self.translate_paragraph, args=(paragraph, translator, target_lang))
                        threads.append(thread)
                        thread.start()
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        thread = threading.Thread(target=self.translate_paragraph, args=(paragraph, translator, target_lang))
                        threads.append(thread)
                        thread.start()

        # Wait for all threads to complete
        for thread in threads:
            thread.join()

    def translate_document(self):
        if not self.input_path or not self.output_path or not self.target_language_code:
            print("Error: Please provide all required paths and target language.")
            return

        try:
            # Load the Word document
            doc = docx.Document(self.input_path)
            translator = Translator()

            # Process paragraphs, tables, headers, and footers
            self.process_paragraphs(doc, translator, self.target_language_code)

            # Save the translated document
            doc.save(self.output_path)
            print(f"Document has been translated and saved as {self.output_path}.")
        except Exception as e:
            print(f"Error: An error occurred: {e}")


# If you want to run this from another script, you can create an instance like this:
# app = WordTranslationApp(input_path="input.docx", output_path="output_translated.docx", target_lang="sr_Latn")
